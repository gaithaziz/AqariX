# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "python-docx",
# ]
# ///

import os
import re
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    md_path = 'docs/ai-priorities.md'
    docx_path = 'docs/ai-priorities.docx'
    
    doc = docx.Document()
    
    # Set standard page margin
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Match headings
        if line.startswith('# '):
            text = line[2:]
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            text = line[3:]
            p = doc.add_heading(text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            text = line[4:]
            p = doc.add_heading(text, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Match bullet points
        elif line.startswith('* ') or line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parse_runs(p, text)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parse_runs(p, line)
            
    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

def parse_runs(paragraph, text):
    # Regex to split bold text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean_part = part[2:-2]
            run = paragraph.add_run(clean_part)
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    main()
